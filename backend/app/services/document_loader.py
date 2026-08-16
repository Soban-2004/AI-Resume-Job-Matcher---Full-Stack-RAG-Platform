from io import BytesIO

import fitz  # PyMuPDF
from docx import Document

# Below this, text is unreadable at any zoom level a human would actually
# use to review a resume -- there's no legitimate reason for real content to
# be this small. This is the classic "ATS keyword stuffing" trick: cramming
# a giant list of skills into 1pt text so it's invisible on screen/print but
# still gets extracted by a naive parser.
_MIN_VISIBLE_FONT_SIZE = 5.0

# Resumes are near-universally laid out on a white base page (colorful
# templates use colored *accent* blocks, not a colored page background), so
# text this close to pure white is either invisible-on-white hidden content
# or, at best, unreadable by a human reviewer either way -- safe to drop.
_NEAR_WHITE_THRESHOLD = 250  # out of 255 per RGB channel


def _rgb_from_int(color: int) -> tuple[int, int, int]:
    return ((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)


def _is_hidden_span(span: dict) -> bool:
    if span.get("size", 12) < _MIN_VISIBLE_FONT_SIZE:
        return True
    r, g, b = _rgb_from_int(span.get("color", 0))
    if r >= _NEAR_WHITE_THRESHOLD and g >= _NEAR_WHITE_THRESHOLD and b >= _NEAR_WHITE_THRESHOLD:
        return True
    return False


def _append_hyperlinks(text: str, links: list[str]) -> str:
    # A resume's GitHub/portfolio/LinkedIn link is very often a clickable
    # word or icon ("GitHub", a logo) rather than visible URL text -- plain
    # text extraction never sees the underlying URL at all in that case,
    # since it lives in a separate link-annotation structure, not the text
    # content itself. Appending the real targets as plain text means
    # url_extractor (and anything else reading resume_text) picks them up
    # for free, with no change needed anywhere else in the pipeline.
    deduped = list(dict.fromkeys(links))
    if not deduped:
        return text
    return text + "\n\n[Linked URLs: " + ", ".join(deduped) + "]"


def _extract_pdf_text(content: bytes) -> str:
    """Extracts visible text only -- filters out spans that are tiny or
    near-white-on-white, the two standard ways invisible keyword stuffing is
    hidden in a PDF. Plain page.get_text() has no concept of "invisible to a
    human" and would extract stuffed content identically to real text.
    """
    pages_text: list[str] = []
    hyperlinks: list[str] = []
    with fitz.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            lines_text: list[str] = []
            for block in page.get_text("dict")["blocks"]:
                for line in block.get("lines", []):
                    visible_spans = [s["text"] for s in line["spans"] if not _is_hidden_span(s)]
                    if visible_spans:
                        lines_text.append("".join(visible_spans))
            pages_text.append("\n".join(lines_text))

            # get_links() reads link *annotations* -- a completely separate
            # structure from the text content stream above, which is why a
            # hyperlinked word never surfaces its URL through get_text() at
            # all, regardless of font size/color filtering.
            for link in page.get_links():
                uri = link.get("uri")
                if uri:
                    hyperlinks.append(uri)

    return _append_hyperlinks("\n".join(pages_text), hyperlinks)


def _extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    # Resumes commonly use tables for column layouts (skills grids, date/role
    # alignment) -- docx2txt included table text by default, so this keeps
    # feature parity with python-docx's paragraph-only default.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    # Hyperlink targets live in the part's relationships, not in any run's
    # .text -- same gap as the PDF case, same fix.
    hyperlink_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    hyperlinks = [
        rel.target_ref for rel in document.part.rels.values() if rel.reltype == hyperlink_rel_type and rel.is_external
    ]

    return _append_hyperlinks("\n".join(parts), hyperlinks)


def load_document(filename: str, content: bytes) -> str | None:
    """Extracts text from an uploaded PDF/DOCX/TXT file, entirely in-memory.

    (The old Streamlit version wrote .docx uploads to a shared temp.docx
    path on disk, which breaks under concurrent requests.)
    """
    try:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            return _extract_pdf_text(content)
        elif lower.endswith(".docx"):
            return _extract_docx_text(content)
        elif lower.endswith(".txt"):
            return content.decode("utf-8", errors="ignore")
        return None
    except Exception:
        return None
